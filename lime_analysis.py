"""
Logistic Regression with LIME Interpretation for Bank/Tabular Data.

- Loads a real dataset from Desktop (prefers files whose name contains
  'real data 1', case-insensitive).
- Prepares data, performs L1-based feature selection.
- Fits Logistic Regression.
- Runs LIME (local explanations) on sample test instances.
- Saves detailed results to:
    - lime_explanations.xlsx
    - lime_feature_importance.xlsx
    - lime_report.docx  (if python-docx is installed)
    - lime_instance_*.html  (interactive LIME HTML for a few cases)

This file is completely independent from the SHAP script.
"""

import os
import re
import sys
import warnings
from pathlib import Path

import matplotlib.pyplot as plt  # noqa: F401  (kept in case of future plots)
import numpy as np
import pandas as pd
from lime.lime_tabular import LimeTabularExplainer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    accuracy_score,
    classification_report,
    confusion_matrix,
    roc_auc_score,
)
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler

warnings.filterwarnings("ignore")

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    Document = None

DESKTOP = Path(os.path.expanduser("~")) / "Desktop"
# All LIME results (Excel, Word, HTML, txt) will be saved under:
#   Desktop/lime.real
OUTPUT_DIR = DESKTOP / "lime.real"
SUPPORTED_EXTENSIONS = [".csv", ".xlsx", ".xls", ".data", ".data-numeric", ".txt"]

# Number of features to keep after selection (most important only)
N_FEATURES_SELECT = 15


def normalize_name_for_match(name: str) -> str:
    """Normalize filename for flexible matching (lowercase, remove spaces/_/-)."""
    return re.sub(r"[\\s_\\-]+", "", name).lower()


def find_data_on_desktop(prefer_real_data_1: bool = True) -> Path | None:
    """
    Search Desktop recursively for a data file.

    If prefer_real_data_1=True, first try to find a file whose name (without
    extension) contains something like 'real data 1' (ignoring spaces/_/- and case).
    If not found, fall back to the largest/most column-rich dataset.
    """
    if not DESKTOP.exists():
        return None

    candidates: list[Path] = []
    for path in DESKTOP.rglob("*"):
        if not path.is_file():
            continue
        if path.suffix.lower() in [".csv", ".xlsx", ".xls"]:
            candidates.append(path)
        elif path.suffix in [".data", ".txt"] or path.name.endswith(".data-numeric"):
            candidates.append(path)

    if not candidates:
        return None

    if prefer_real_data_1:
        target_key = normalize_name_for_match("real data 1")
        prioritized: list[Path] = []
        for p in candidates:
            base = p.stem  # filename without extension
            if target_key in normalize_name_for_match(base):
                prioritized.append(p)
        if prioritized:
            # If multiple match, pick the one with most columns/rows
            def cols_rows(q: Path) -> tuple[int, int]:
                try:
                    if q.suffix.lower() in [".xlsx", ".xls"]:
                        d = pd.read_excel(q)
                    elif q.suffix.lower() == ".csv":
                        d = pd.read_csv(q)
                    else:
                        d = pd.read_csv(q, sep=r"\\s+", header=None)
                    return d.shape[1], d.shape[0]
                except Exception:
                    return 0, 0

            prioritized.sort(key=cols_rows, reverse=True)
            return prioritized[0]

    # Fallback: choose "largest" dataset (by columns, then rows)
    def cols_rows(p: Path) -> tuple[int, int]:
        try:
            if p.suffix.lower() in [".xlsx", ".xls"]:
                d = pd.read_excel(p)
            elif p.suffix.lower() == ".csv":
                d = pd.read_csv(p)
            else:
                d = pd.read_csv(p, sep=r"\\s+", header=None)
            return d.shape[1], d.shape[0]
        except Exception:
            return 0, 0

    candidates.sort(key=cols_rows, reverse=True)
    return candidates[0]


def load_german_data_numeric(path: Path) -> pd.DataFrame:
    df = pd.read_csv(path, sep=r"\\s+", header=None)
    df.columns = [f"X{i}" for i in range(len(df.columns) - 1)] + ["target"]
    df["target"] = df["target"].replace({2: 0})
    return df


def load_german_data_raw(path: Path) -> pd.DataFrame:
    df = pd.read_csv(path, sep=r"\\s+", header=None)
    col_names = [
        "checking_status",
        "duration",
        "credit_history",
        "purpose",
        "credit_amount",
        "savings",
        "employment",
        "installment_rate",
        "personal_status",
        "other_parties",
        "residence_since",
        "property_magnitude",
        "age",
        "other_plans",
        "housing",
        "num_credits",
        "job",
        "num_dependents",
        "phone",
        "foreign_worker",
        "target",
    ]
    if len(df.columns) == 21:
        df.columns = col_names
    else:
        df.columns = [f"X{i}" for i in range(len(df.columns) - 1)] + ["target"]
    df["target"] = df["target"].replace({2: 0})
    for col in df.select_dtypes(include=["object"]).columns:
        df[col] = pd.Categorical(df[col]).codes
    return df


def load_csv_or_excel(path: Path) -> pd.DataFrame:
    if path.suffix.lower() == ".csv":
        return pd.read_csv(path)
    return pd.read_excel(path)


def load_data(path: Path) -> pd.DataFrame:
    name_lower = path.name.lower()
    if "german" in name_lower and ("numeric" in name_lower or path.suffix == ".data-numeric"):
        return load_german_data_numeric(path)
    if "german" in name_lower and path.suffix == ".data":
        return load_german_data_raw(path)
    return load_csv_or_excel(path)


def prepare_data(df: pd.DataFrame, target_col: str | None = None) -> tuple[pd.DataFrame, pd.Series, list[str]]:
    df = df.dropna()
    if target_col and target_col in df.columns:
        y = df[target_col]
        X = df.drop(columns=[target_col])
    else:
        for name in ["target", "Target", "class", "outcome", "y", "label"]:
            if name in df.columns:
                target_col = name
                break
        if target_col is None:
            target_col = df.columns[-1]
        y = df[target_col]
        X = df.drop(columns=[target_col])
    X = X.select_dtypes(include=[np.number])
    if X.empty:
        raise ValueError("No numeric features found. Encode categorical columns first.")
    feature_names = X.columns.tolist()
    return X, y, feature_names


def select_features_l1(
    X_train: np.ndarray, y_train, feature_names: list[str], n_select: int
) -> tuple[np.ndarray, np.ndarray, list[str], list[int], StandardScaler]:
    """
    Select top n_select features using L1 (Lasso) logistic regression.

    Returns:
        X_train_sel:   subset of X_train with selected features (unscaled)
        X_scaled_sel:  scaled selected features
        selected_names: list of feature names
        selected_idx:   list of integer indices into original feature list
        scaler:         fitted StandardScaler on selected features
    """
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X_train)
    model_l1 = LogisticRegression(
        penalty="l1",
        solver="saga",
        max_iter=2000,
        random_state=42,
        C=0.5,
    )
    model_l1.fit(X_scaled, y_train)
    coef = np.abs(model_l1.coef_).ravel()
    top_idx = np.argsort(coef)[-n_select:][::-1]
    selected_names = [feature_names[i] for i in top_idx]
    X_train_sel = X_train[:, top_idx]

    scaler_selected = StandardScaler()
    X_scaled_sel = scaler_selected.fit_transform(X_train_sel)
    return X_train_sel, X_scaled_sel, selected_names, top_idx.tolist(), scaler_selected


def fit_logistic(
    X_train_scaled: np.ndarray, y_train, X_test_scaled: np.ndarray, y_test
) -> LogisticRegression:
    model = LogisticRegression(
        penalty="l2",
        solver="lbfgs",
        max_iter=2000,
        random_state=42,
        C=1.0,
    )
    model.fit(X_train_scaled, y_train)
    y_pred = model.predict(X_test_scaled)
    acc = accuracy_score(y_test, y_pred)
    try:
        auc = roc_auc_score(y_test, model.predict_proba(X_test_scaled)[:, 1])
        print(f"  Accuracy: {acc:.4f}  |  AUC-ROC: {auc:.4f}")
    except Exception:
        print(f"  Accuracy: {acc:.4f}")
    return model


def run_lime_for_logistic(
    model: LogisticRegression,
    X_train_scaled: np.ndarray,
    X_test_scaled: np.ndarray,
    y_train: pd.Series,
    y_test: pd.Series,
    feature_names: list[str],
    output_dir: Path,
    n_instances: int = 10,
) -> None:
    """
    Run LIME on a subset of test instances, save:
      - detailed local explanations to Excel
      - aggregated "global" importance to Excel
      - a short Word report (if python-docx is available)
      - HTML explanations for first few instances
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    class_labels = sorted(pd.unique(y_train))
    class_names = [str(c) for c in class_labels]

    explainer = LimeTabularExplainer(
        training_data=X_train_scaled,
        feature_names=feature_names,
        class_names=class_names,
        mode="classification",
        discretize_continuous=True,
    )

    n_instances = min(n_instances, X_test_scaled.shape[0])
    rows: list[dict] = []

    print(f"\nRunning LIME for {n_instances} test instances...")
    for i in range(n_instances):
        x_row = X_test_scaled[i]
        true_label = y_test.iloc[i]
        proba = model.predict_proba(x_row.reshape(1, -1))[0]
        pred_class_idx = int(np.argmax(proba))
        pred_class = class_labels[pred_class_idx]
        pred_proba = float(proba[pred_class_idx])

        exp = explainer.explain_instance(
            data_row=x_row,
            predict_fn=model.predict_proba,
            num_features=min(10, len(feature_names)),
        )

        # Save interactive HTML for first few instances
        if i < 5:
            html_path = output_dir / f"lime_instance_{i+1}.html"
            try:
                exp.save_to_file(str(html_path))
                print(f"  Saved HTML explanation: {html_path.name}")
            except Exception as e:
                print(f"  Could not save HTML for instance {i+1}: {e}")

        for feat, weight in exp.as_list():
            rows.append(
                {
                    "instance_index": i,
                    "true_label": true_label,
                    "predicted_label": pred_class,
                    "predicted_proba_for_predicted_label": pred_proba,
                    "feature": feat,
                    "weight": float(weight),
                    "weight_abs": float(abs(weight)),
                }
            )

    if not rows:
        print("No LIME explanations could be generated.")
        return

    explanations_df = pd.DataFrame(rows)
    explanations_path = output_dir / "lime_explanations.xlsx"
    try:
        explanations_df.to_excel(explanations_path, index=False)
        print(f"  Saved: {explanations_path.name}")
    except Exception as e:
        print(f"  Could not save Excel {explanations_path.name}: {e}")

    global_imp = (
        explanations_df.groupby("feature")["weight_abs"]
        .mean()
        .reset_index()
        .sort_values("weight_abs", ascending=False)
    )
    global_imp_path = output_dir / "lime_feature_importance.xlsx"
    try:
        global_imp.to_excel(global_imp_path, index=False)
        print(f"  Saved: {global_imp_path.name}")
    except Exception as e:
        print(f"  Could not save Excel {global_imp_path.name}: {e}")

    if Document is not None:
        try:
            doc = Document()
            doc.add_heading("Logistic Regression with LIME Interpretation", level=1)
            doc.add_paragraph(
                "This report summarizes LIME-based local explanations for a "
                "logistic regression model fitted on the selected real dataset."
            )

            doc.add_heading("Model and Data", level=2)
            doc.add_paragraph(f"Number of training instances: {X_train_scaled.shape[0]}")
            doc.add_paragraph(f"Number of test instances: {X_test_scaled.shape[0]}")
            doc.add_paragraph(f"Number of selected features: {len(feature_names)}")

            doc.add_heading("Top Features by Average LIME Weight", level=2)
            table = doc.add_table(
                rows=min(11, len(global_imp) + 1), cols=2
            )  # header + top 10
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Feature"
            hdr_cells[1].text = "Mean |weight|"
            for idx, (_, row) in enumerate(global_imp.head(10).iterrows(), start=1):
                table.rows[idx].cells[0].text = str(row["feature"])
                table.rows[idx].cells[1].text = f"{row['weight_abs']:.4f}"

            doc.add_paragraph(
                "For more detailed, instance-level explanations see the accompanying "
                "Excel file 'lime_explanations.xlsx' and HTML files "
                "'lime_instance_*.html'."
            )

            report_path = output_dir / "lime_report.docx"
            doc.save(report_path)
            print(f"  Saved: {report_path.name}")
        except Exception as e:
            print(f"  Could not create Word report: {e}")
    else:
        print(
            "python-docx not installed: skipping Word report. "
            "Install with 'pip install python-docx' if you need it."
        )


def main() -> None:
    print("=" * 60)
    print("Logistic Regression with LIME (feature selection applied)")
    print("=" * 60)

    data_path = find_data_on_desktop(prefer_real_data_1=True)
    if data_path is None:
        print("\nNo data file found on Desktop.")
        manual = input("Enter full path to data file (or Enter to exit): ").strip()
        if not manual:
            sys.exit(1)
        data_path = Path(manual)

    if not data_path.exists():
        print(f"File not found: {data_path}")
        sys.exit(1)

    try:
        print(f"\nData file (selected for LIME): {data_path}")
    except UnicodeEncodeError:
        print("\nData file (selected for LIME): [path with non-ASCII characters]")

    try:
        df = load_data(data_path)
        print(f"Loaded: {len(df)} rows, {len(df.columns)} columns")
    except Exception as e:
        print(f"Error loading data: {e}")
        sys.exit(1)

    try:
        X, y, feature_names = prepare_data(df)
        print(f"Initial features: {len(feature_names)} | Target: {y.name}")
        print(f"Target distribution:\n{y.value_counts()}")
    except Exception as e:
        print(f"Error preparing data: {e}")
        sys.exit(1)

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=42
    )
    n_select = min(N_FEATURES_SELECT, X_train.shape[1])
    print(f"\nFeature selection: keeping top {n_select} features (L1 logistic regression)")

    (
        X_train_sel,
        X_train_scaled_sel,
        selected_names,
        selected_idx,
        scaler_selected,
    ) = select_features_l1(X_train.values, y_train, feature_names, n_select)

    X_test_sel = X_test.values[:, selected_idx]
    X_test_scaled_sel = scaler_selected.transform(X_test_sel)

    print(f"Selected features: {selected_names}")

    print("\nFitting logistic regression (L2) on selected features...")
    model = fit_logistic(X_train_scaled_sel, y_train, X_test_scaled_sel, y_test)

    print("\nRunning LIME analysis on test set...")
    run_lime_for_logistic(
        model=model,
        X_train_scaled=X_train_scaled_sel,
        X_test_scaled=X_test_scaled_sel,
        y_train=y_train,
        y_test=y_test,
        feature_names=selected_names,
        output_dir=OUTPUT_DIR,
        n_instances=10,
    )

    # Save model summary and selected feature list for reference / reporting
    summary_path = OUTPUT_DIR / "model_summary.txt"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    try:
        with open(summary_path, "w", encoding="utf-8") as f:
            f.write("Model: Logistic Regression (L2)\n")
            f.write(f"Features selected: {n_select}\n")
            f.write(f"Selected feature names:\n{selected_names}\n\n")
            f.write("Classification report:\n")
            f.write(classification_report(y_test, model.predict(X_test_scaled_sel)))
            f.write("\nConfusion matrix:\n")
            f.write(str(confusion_matrix(y_test, model.predict(X_test_scaled_sel))))
        print(f"  Saved: {summary_path.name}")
    except Exception as e:
        print(f"  Could not save model summary: {e}")

    print("\n" + "=" * 60)
    print(f"Done. LIME results saved to: {OUTPUT_DIR}")
    print("=" * 60)


if __name__ == "__main__":
    main()

